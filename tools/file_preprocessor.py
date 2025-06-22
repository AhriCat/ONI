import os
import logging
from typing import Dict, List, Optional, Any, Union, BinaryIO, TextIO
import mimetypes
import tempfile
import shutil

# Optional imports - will be loaded dynamically when needed
_LIBRARIES = {
    'pdf': 'PyPDF2',
    'docx': 'docx',
    'image': 'PIL',
    'excel': 'pandas',
    'csv': 'pandas',
    'json': 'json',
    'xml': 'xml.etree.ElementTree',
    'html': 'bs4',
    'audio': 'librosa',
    'video': 'cv2',
    'code': 'pygments'
}

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class FilePreprocessor:
    """
    A modular file preprocessor that handles various file types and allows ONI
    to choose which components to feed files to.
    """
    def __init__(self):
        """Initialize the file preprocessor."""
        self.temp_dir = tempfile.mkdtemp(prefix="oni_file_processor_")
        self.loaded_libraries = {}
        self.registered_handlers = {}
        self.register_default_handlers()
        
    def __del__(self):
        """Clean up temporary directory when the object is destroyed."""
        try:
            shutil.rmtree(self.temp_dir)
        except Exception as e:
            logger.error(f"Error cleaning up temporary directory: {e}")
    
    def _load_library(self, library_key: str) -> bool:
        """
        Dynamically load a library when needed.
        
        Args:
            library_key: Key of the library to load
            
        Returns:
            bool: True if library was loaded successfully, False otherwise
        """
        if library_key in self.loaded_libraries:
            return True
            
        if library_key not in _LIBRARIES:
            logger.error(f"Unknown library key: {library_key}")
            return False
            
        library_name = _LIBRARIES[library_key]
        
        try:
            if library_key == 'pdf':
                import PyPDF2
                self.loaded_libraries[library_key] = PyPDF2
            elif library_key == 'docx':
                import docx
                self.loaded_libraries[library_key] = docx
            elif library_key == 'image':
                from PIL import Image
                self.loaded_libraries[library_key] = Image
            elif library_key == 'excel' or library_key == 'csv':
                import pandas as pd
                self.loaded_libraries[library_key] = pd
            elif library_key == 'json':
                import json
                self.loaded_libraries[library_key] = json
            elif library_key == 'xml':
                import xml.etree.ElementTree as ET
                self.loaded_libraries[library_key] = ET
            elif library_key == 'html':
                from bs4 import BeautifulSoup
                self.loaded_libraries[library_key] = BeautifulSoup
            elif library_key == 'audio':
                import librosa
                self.loaded_libraries[library_key] = librosa
            elif library_key == 'video':
                import cv2
                self.loaded_libraries[library_key] = cv2
            elif library_key == 'code':
                import pygments
                from pygments import lexers, formatters
                self.loaded_libraries[library_key] = pygments
            
            logger.info(f"Successfully loaded library: {library_name}")
            return True
        except ImportError:
            logger.warning(f"Could not import {library_name}. Some functionality may be limited.")
            return False
    
    def register_default_handlers(self):
        """Register default file handlers."""
        # Text files
        self.register_handler('.txt', self._process_text_file)
        self.register_handler('.md', self._process_text_file)
        self.register_handler('.csv', self._process_csv_file)
        self.register_handler('.json', self._process_json_file)
        self.register_handler('.xml', self._process_xml_file)
        self.register_handler('.html', self._process_html_file)
        
        # Document files
        self.register_handler('.pdf', self._process_pdf_file)
        self.register_handler('.docx', self._process_docx_file)
        self.register_handler('.xlsx', self._process_excel_file)
        self.register_handler('.xls', self._process_excel_file)
        
        # Image files
        self.register_handler('.jpg', self._process_image_file)
        self.register_handler('.jpeg', self._process_image_file)
        self.register_handler('.png', self._process_image_file)
        self.register_handler('.gif', self._process_image_file)
        
        # Audio files
        self.register_handler('.mp3', self._process_audio_file)
        self.register_handler('.wav', self._process_audio_file)
        self.register_handler('.ogg', self._process_audio_file)
        
        # Video files
        self.register_handler('.mp4', self._process_video_file)
        self.register_handler('.avi', self._process_video_file)
        self.register_handler('.mov', self._process_video_file)
        
        # Code files
        self.register_handler('.py', self._process_code_file)
        self.register_handler('.js', self._process_code_file)
        self.register_handler('.java', self._process_code_file)
        self.register_handler('.cpp', self._process_code_file)
        self.register_handler('.c', self._process_code_file)
        self.register_handler('.h', self._process_code_file)
        self.register_handler('.cs', self._process_code_file)
        self.register_handler('.go', self._process_code_file)
        self.register_handler('.rb', self._process_code_file)
        self.register_handler('.php', self._process_code_file)
    
    def register_handler(self, extension: str, handler_func):
        """
        Register a custom file handler.
        
        Args:
            extension: File extension (including the dot)
            handler_func: Function to handle the file
        """
        self.registered_handlers[extension.lower()] = handler_func
        logger.info(f"Registered handler for {extension} files")
    
    def process_file(self, file_path: str, target_format: str = 'text', 
                    options: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Process a file and return its content in the requested format.
        
        Args:
            file_path: Path to the file
            target_format: Desired output format ('text', 'binary', 'tensor', etc.)
            options: Additional processing options
            
        Returns:
            Dict containing processed content and metadata
        """
        if not os.path.exists(file_path):
            return {
                'success': False,
                'error': f"File not found: {file_path}",
                'file_path': file_path
            }
        
        options = options or {}
        
        try:
            # Get file extension
            _, extension = os.path.splitext(file_path)
            extension = extension.lower()
            
            # Get file size
            file_size = os.path.getsize(file_path)
            
            # Get file mime type
            mime_type, _ = mimetypes.guess_type(file_path)
            
            # Check if we have a registered handler for this extension
            if extension in self.registered_handlers:
                handler = self.registered_handlers[extension]
                result = handler(file_path, target_format, options)
                
                # Add file metadata
                result.update({
                    'file_path': file_path,
                    'file_name': os.path.basename(file_path),
                    'file_size': file_size,
                    'mime_type': mime_type,
                    'extension': extension
                })
                
                return result
            else:
                # Try to process as binary file
                return self._process_binary_file(file_path, target_format, options)
                
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_path': file_path
            }
    
    def process_files(self, file_paths: List[str], target_format: str = 'text',
                     options: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """
        Process multiple files.
        
        Args:
            file_paths: List of file paths
            target_format: Desired output format
            options: Additional processing options
            
        Returns:
            List of processing results
        """
        results = []
        for file_path in file_paths:
            result = self.process_file(file_path, target_format, options)
            results.append(result)
        return results
    
    def _process_text_file(self, file_path: str, target_format: str, 
                          options: Dict[str, Any]) -> Dict[str, Any]:
        """Process a text file."""
        encoding = options.get('encoding', 'utf-8')
        max_length = options.get('max_length', None)
        
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
                
            if max_length and len(content) > max_length:
                content = content[:max_length]
                
            return {
                'success': True,
                'content': content,
                'content_type': 'text',
                'encoding': encoding
            }
        except UnicodeDecodeError:
            # Try with a different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                    
                if max_length and len(content) > max_length:
                    content = content[:max_length]
                    
                return {
                    'success': True,
                    'content': content,
                    'content_type': 'text',
                    'encoding': 'latin-1'
                }
            except Exception as e:
                return {
                    'success': False,
                    'error': f"Failed to decode text file: {e}",
                    'content_type': 'text'
                }
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing text file: {e}",
                'content_type': 'text'
            }
    
    def _process_pdf_file(self, file_path: str, target_format: str, 
                         options: Dict[str, Any]) -> Dict[str, Any]:
        """Process a PDF file."""
        if not self._load_library('pdf'):
            return {
                'success': False,
                'error': "PyPDF2 library not available",
                'content_type': 'pdf'
            }
            
        PyPDF2 = self.loaded_libraries['pdf']
        max_pages = options.get('max_pages', None)
        extract_images = options.get('extract_images', False)
        
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                num_pages = len(reader.pages)
                
                if max_pages:
                    num_pages = min(num_pages, max_pages)
                
                text_content = ""
                for i in range(num_pages):
                    page = reader.pages[i]
                    text_content += page.extract_text() + "\n\n"
                
                result = {
                    'success': True,
                    'content': text_content,
                    'content_type': 'pdf',
                    'num_pages': num_pages,
                    'total_pages': len(reader.pages)
                }
                
                # Extract images if requested
                if extract_images and self._load_library('image'):
                    Image = self.loaded_libraries['image']
                    images = []
                    
                    for i in range(num_pages):
                        page = reader.pages[i]
                        if '/XObject' in page['/Resources']:
                            xObject = page['/Resources']['/XObject']
                            for obj in xObject:
                                if xObject[obj]['/Subtype'] == '/Image':
                                    try:
                                        data = xObject[obj].get_data()
                                        img_path = os.path.join(self.temp_dir, f"image_{i}_{obj}.png")
                                        with open(img_path, 'wb') as img_file:
                                            img_file.write(data)
                                        images.append(img_path)
                                    except:
                                        pass
                    
                    result['images'] = images
                    result['num_images'] = len(images)
                
                return result
                
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing PDF file: {e}",
                'content_type': 'pdf'
            }
    
    def _process_docx_file(self, file_path: str, target_format: str, 
                          options: Dict[str, Any]) -> Dict[str, Any]:
        """Process a DOCX file."""
        if not self._load_library('docx'):
            return {
                'success': False,
                'error': "python-docx library not available",
                'content_type': 'docx'
            }
            
        docx = self.loaded_libraries['docx']
        extract_images = options.get('extract_images', False)
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            result = {
                'success': True,
                'content': '\n\n'.join(paragraphs),
                'content_type': 'docx',
                'paragraphs': paragraphs,
                'tables': tables,
                'num_paragraphs': len(paragraphs),
                'num_tables': len(tables)
            }
            
            # Extract images if requested
            if extract_images:
                images = []
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_data = rel.target_part.blob
                            img_path = os.path.join(self.temp_dir, f"image_{rel.rId}.png")
                            with open(img_path, 'wb') as img_file:
                                img_file.write(image_data)
                            images.append(img_path)
                        except:
                            pass
                
                result['images'] = images
                result['num_images'] = len(images)
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing DOCX file: {e}",
                'content_type': 'docx'
            }
    
    def _process_image_file(self, file_path: str, target_format: str, 
                           options: Dict[str, Any]) -> Dict[str, Any]:
        """Process an image file."""
        if not self._load_library('image'):
            return {
                'success': False,
                'error': "PIL library not available",
                'content_type': 'image'
            }
            
        Image = self.loaded_libraries['image']
        resize = options.get('resize', None)
        convert_to_grayscale = options.get('grayscale', False)
        extract_text = options.get('extract_text', False)
        
        try:
            with Image.open(file_path) as img:
                # Get image info
                width, height = img.size
                format_name = img.format
                mode = img.mode
                
                # Resize if requested
                if resize:
                    img = img.resize(resize)
                    width, height = img.size
                
                # Convert to grayscale if requested
                if convert_to_grayscale:
                    img = img.convert('L')
                    mode = img.mode
                
                # Save processed image
                processed_path = os.path.join(self.temp_dir, f"processed_{os.path.basename(file_path)}")
                img.save(processed_path)
                
                result = {
                    'success': True,
                    'content': processed_path,
                    'content_type': 'image',
                    'width': width,
                    'height': height,
                    'format': format_name,
                    'mode': mode
                }
                
                # Extract text if requested
                if extract_text:
                    if self._load_library('image') and hasattr(self.loaded_libraries['image'], 'pytesseract'):
                        pytesseract = self.loaded_libraries['image'].pytesseract
                        text = pytesseract.image_to_string(img)
                        result['extracted_text'] = text
                    else:
                        result['warning'] = "pytesseract not available for text extraction"
                
                return result
                
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing image file: {e}",
                'content_type': 'image'
            }
    
    def _process_audio_file(self, file_path: str, target_format: str, 
                           options: Dict[str, Any]) -> Dict[str, Any]:
        """Process an audio file."""
        if not self._load_library('audio'):
            return {
                'success': False,
                'error': "librosa library not available",
                'content_type': 'audio'
            }
            
        librosa = self.loaded_libraries['audio']
        duration = options.get('duration', None)
        sample_rate = options.get('sample_rate', 22050)
        
        try:
            # Load audio file
            y, sr = librosa.load(file_path, sr=sample_rate, duration=duration)
            
            result = {
                'success': True,
                'content': y.tolist() if target_format == 'list' else y,
                'content_type': 'audio',
                'sample_rate': sr,
                'duration': librosa.get_duration(y=y, sr=sr),
                'num_samples': len(y)
            }
            
            # Extract features if requested
            if options.get('extract_features', False):
                # Extract mel spectrogram
                mel_spec = librosa.feature.melspectrogram(y=y, sr=sr)
                result['mel_spectrogram'] = mel_spec.tolist() if target_format == 'list' else mel_spec
                
                # Extract MFCC
                mfcc = librosa.feature.mfcc(y=y, sr=sr, n_mfcc=13)
                result['mfcc'] = mfcc.tolist() if target_format == 'list' else mfcc
                
                # Extract chroma
                chroma = librosa.feature.chroma_stft(y=y, sr=sr)
                result['chroma'] = chroma.tolist() if target_format == 'list' else chroma
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing audio file: {e}",
                'content_type': 'audio'
            }
    
    def _process_video_file(self, file_path: str, target_format: str, 
                           options: Dict[str, Any]) -> Dict[str, Any]:
        """Process a video file."""
        if not self._load_library('video'):
            return {
                'success': False,
                'error': "OpenCV library not available",
                'content_type': 'video'
            }
            
        cv2 = self.loaded_libraries['video']
        max_frames = options.get('max_frames', 10)
        frame_interval = options.get('frame_interval', 30)
        resize = options.get('resize', None)
        
        try:
            # Open video file
            cap = cv2.VideoCapture(file_path)
            
            # Get video properties
            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            duration = frame_count / fps if fps > 0 else 0
            
            # Extract frames
            frames = []
            frame_paths = []
            
            for i in range(min(max_frames, frame_count)):
                # Set frame position
                cap.set(cv2.CAP_PROP_POS_FRAMES, i * frame_interval)
                
                # Read frame
                ret, frame = cap.read()
                if not ret:
                    break
                
                # Resize if requested
                if resize:
                    frame = cv2.resize(frame, resize)
                
                # Save frame
                frame_path = os.path.join(self.temp_dir, f"frame_{i}.jpg")
                cv2.imwrite(frame_path, frame)
                frame_paths.append(frame_path)
                
                if target_format == 'tensor' or target_format == 'list':
                    frames.append(frame)
            
            # Release video capture
            cap.release()
            
            result = {
                'success': True,
                'content_type': 'video',
                'fps': fps,
                'frame_count': frame_count,
                'width': width,
                'height': height,
                'duration': duration,
                'frame_paths': frame_paths,
                'num_extracted_frames': len(frame_paths)
            }
            
            if target_format == 'tensor' or target_format == 'list':
                result['frames'] = frames if target_format == 'tensor' else [f.tolist() for f in frames]
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing video file: {e}",
                'content_type': 'video'
            }
    
    def _process_csv_file(self, file_path: str, target_format: str, 
                         options: Dict[str, Any]) -> Dict[str, Any]:
        """Process a CSV file."""
        if not self._load_library('csv'):
            return {
                'success': False,
                'error': "pandas library not available",
                'content_type': 'csv'
            }
            
        pd = self.loaded_libraries['csv']
        max_rows = options.get('max_rows', None)
        
        try:
            # Read CSV file
            df = pd.read_csv(file_path)
            
            # Limit rows if requested
            if max_rows:
                df = df.head(max_rows)
            
            result = {
                'success': True,
                'content_type': 'csv',
                'num_rows': len(df),
                'num_columns': len(df.columns),
                'columns': df.columns.tolist(),
                'dtypes': {col: str(dtype) for col, dtype in df.dtypes.items()}
            }
            
            # Return data in requested format
            if target_format == 'text':
                result['content'] = df.to_string()
            elif target_format == 'json':
                result['content'] = df.to_json(orient='records')
            elif target_format == 'list':
                result['content'] = df.values.tolist()
                result['headers'] = df.columns.tolist()
            elif target_format == 'dict':
                result['content'] = df.to_dict(orient='records')
            else:
                result['content'] = df
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing CSV file: {e}",
                'content_type': 'csv'
            }
    
    def _process_excel_file(self, file_path: str, target_format: str, 
                           options: Dict[str, Any]) -> Dict[str, Any]:
        """Process an Excel file."""
        if not self._load_library('excel'):
            return {
                'success': False,
                'error': "pandas library not available",
                'content_type': 'excel'
            }
            
        pd = self.loaded_libraries['excel']
        sheet_name = options.get('sheet_name', 0)
        max_rows = options.get('max_rows', None)
        
        try:
            # Read Excel file
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # Get sheet names
            xl = pd.ExcelFile(file_path)
            sheet_names = xl.sheet_names
            
            # Limit rows if requested
            if max_rows:
                df = df.head(max_rows)
            
            result = {
                'success': True,
                'content_type': 'excel',
                'num_rows': len(df),
                'num_columns': len(df.columns),
                'columns': df.columns.tolist(),
                'sheet_names': sheet_names,
                'current_sheet': sheet_name,
                'dtypes': {col: str(dtype) for col, dtype in df.dtypes.items()}
            }
            
            # Return data in requested format
            if target_format == 'text':
                result['content'] = df.to_string()
            elif target_format == 'json':
                result['content'] = df.to_json(orient='records')
            elif target_format == 'list':
                result['content'] = df.values.tolist()
                result['headers'] = df.columns.tolist()
            elif target_format == 'dict':
                result['content'] = df.to_dict(orient='records')
            else:
                result['content'] = df
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing Excel file: {e}",
                'content_type': 'excel'
            }
    
    def _process_json_file(self, file_path: str, target_format: str, 
                          options: Dict[str, Any]) -> Dict[str, Any]:
        """Process a JSON file."""
        if not self._load_library('json'):
            return {
                'success': False,
                'error': "json library not available",
                'content_type': 'json'
            }
            
        json = self.loaded_libraries['json']
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            result = {
                'success': True,
                'content': data,
                'content_type': 'json'
            }
            
            # Return data in requested format
            if target_format == 'text':
                result['content'] = json.dumps(data, indent=2)
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing JSON file: {e}",
                'content_type': 'json'
            }
    
    def _process_xml_file(self, file_path: str, target_format: str, 
                         options: Dict[str, Any]) -> Dict[str, Any]:
        """Process an XML file."""
        if not self._load_library('xml'):
            return {
                'success': False,
                'error': "xml.etree.ElementTree library not available",
                'content_type': 'xml'
            }
            
        ET = self.loaded_libraries['xml']
        
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Convert XML to dictionary
            def xml_to_dict(element):
                result = {}
                for child in element:
                    child_data = xml_to_dict(child)
                    if child.tag in result:
                        if type(result[child.tag]) is list:
                            result[child.tag].append(child_data)
                        else:
                            result[child.tag] = [result[child.tag], child_data]
                    else:
                        result[child.tag] = child_data
                
                if element.text and element.text.strip():
                    if not result:
                        return element.text.strip()
                    else:
                        result['_text'] = element.text.strip()
                
                if element.attrib:
                    result['_attributes'] = element.attrib
                
                return result
            
            xml_dict = {root.tag: xml_to_dict(root)}
            
            result = {
                'success': True,
                'content': xml_dict,
                'content_type': 'xml',
                'root_tag': root.tag
            }
            
            # Return data in requested format
            if target_format == 'text':
                with open(file_path, 'r', encoding='utf-8') as f:
                    result['content'] = f.read()
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing XML file: {e}",
                'content_type': 'xml'
            }
    
    def _process_html_file(self, file_path: str, target_format: str, 
                          options: Dict[str, Any]) -> Dict[str, Any]:
        """Process an HTML file."""
        if not self._load_library('html'):
            return {
                'success': False,
                'error': "BeautifulSoup library not available",
                'content_type': 'html'
            }
            
        BeautifulSoup = self.loaded_libraries['html']
        extract_links = options.get('extract_links', False)
        extract_images = options.get('extract_images', False)
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract text
            text = soup.get_text(separator='\n', strip=True)
            
            result = {
                'success': True,
                'content': text if target_format == 'text' else html_content,
                'content_type': 'html',
                'title': soup.title.string if soup.title else None
            }
            
            # Extract links if requested
            if extract_links:
                links = []
                for link in soup.find_all('a'):
                    href = link.get('href')
                    if href:
                        links.append({
                            'text': link.text.strip(),
                            'href': href
                        })
                result['links'] = links
            
            # Extract images if requested
            if extract_images:
                images = []
                for img in soup.find_all('img'):
                    src = img.get('src')
                    if src:
                        images.append({
                            'src': src,
                            'alt': img.get('alt', ''),
                            'width': img.get('width', ''),
                            'height': img.get('height', '')
                        })
                result['images'] = images
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing HTML file: {e}",
                'content_type': 'html'
            }
    
    def _process_code_file(self, file_path: str, target_format: str, 
                          options: Dict[str, Any]) -> Dict[str, Any]:
        """Process a code file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                code = f.read()
            
            result = {
                'success': True,
                'content': code,
                'content_type': 'code',
                'language': os.path.splitext(file_path)[1][1:]  # Get language from extension
            }
            
            # Syntax highlighting if requested and available
            if options.get('syntax_highlight', False) and self._load_library('code'):
                pygments = self.loaded_libraries['code']
                lexer = pygments.lexers.get_lexer_for_filename(file_path)
                formatter = pygments.formatters.HtmlFormatter()
                highlighted_code = pygments.highlight(code, lexer, formatter)
                result['highlighted_code'] = highlighted_code
                result['highlight_css'] = formatter.get_style_defs('.highlight')
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing code file: {e}",
                'content_type': 'code'
            }
    
    def _process_binary_file(self, file_path: str, target_format: str, 
                            options: Dict[str, Any]) -> Dict[str, Any]:
        """Process a binary file."""
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            
            result = {
                'success': True,
                'content_type': 'binary',
                'size': len(content)
            }
            
            # Return data in requested format
            if target_format == 'binary':
                result['content'] = content
            elif target_format == 'hex':
                result['content'] = content.hex()
            elif target_format == 'base64':
                import base64
                result['content'] = base64.b64encode(content).decode('ascii')
            else:
                result['content'] = f"Binary data ({len(content)} bytes)"
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing binary file: {e}",
                'content_type': 'binary'
            }
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """
        Get information about a file without processing its contents.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dict containing file information
        """
        if not os.path.exists(file_path):
            return {
                'success': False,
                'error': f"File not found: {file_path}",
                'file_path': file_path
            }
        
        try:
            # Get file stats
            stats = os.stat(file_path)
            
            # Get file extension and mime type
            _, extension = os.path.splitext(file_path)
            mime_type, _ = mimetypes.guess_type(file_path)
            
            return {
                'success': True,
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'extension': extension.lower(),
                'mime_type': mime_type,
                'size': stats.st_size,
                'created': stats.st_ctime,
                'modified': stats.st_mtime,
                'accessed': stats.st_atime,
                'is_binary': not self._is_text_file(file_path)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Error getting file info: {e}",
                'file_path': file_path
            }
    
    def _is_text_file(self, file_path: str) -> bool:
        """
        Check if a file is a text file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            bool: True if the file is a text file, False otherwise
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                f.read(1024)  # Try to read some content
            return True
        except UnicodeDecodeError:
            return False
        except Exception:
            return False

# Example usage
if __name__ == "__main__":
    processor = FilePreprocessor()
    
    # Process a text file
    result = processor.process_file("example.txt")
    print(result)
    
    # Process a PDF file
    result = processor.process_file("example.pdf", options={'extract_images': True})
    print(result)
    
    # Process an image file
    result = processor.process_file("example.jpg", options={'resize': (800, 600), 'extract_text': True})
    print(result)